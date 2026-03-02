import streamlit as st
import google.generativeai as genai
import json
import time
import random
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from docx import Document
from pypdf import PdfReader
import io

# --- 1. CẤU HÌNH HỆ THỐNG ---
st.set_page_config(
    page_title="TRỢ LÝ PCCC CHUYÊN DỤNG",
    page_icon="🔥",
    layout="centered", 
    initial_sidebar_state="collapsed", 
    menu_items={
        'Get Help': None,
        'Report a bug': None,
        'About': "Hệ thống hỗ trợ nghiệp vụ PCCC PC07 - Phát triển bởi Đại úy Phạm Tùng Linh"
    }
)

# --- 2. CSS "CLEAN & BEAUTIFUL" (ẨN TOÀN BỘ RÁC) ---
st.markdown("""
<style>
    /* 1. ẨN CÁC THÀNH PHẦN MẶC ĐỊNH CỦA STREAMLIT */
    #MainMenu {visibility: hidden;} /* Ẩn menu 3 gạch góc phải */
    footer {visibility: hidden;} /* Ẩn footer 'Made with Streamlit' */
    header {visibility: hidden;} /* Ẩn thanh header rỗng ở trên cùng */
    .stDeployButton {display:none;} /* Ẩn nút Deploy */
    [data-testid="stToolbar"] {visibility: hidden;} /* Ẩn toolbar hệ thống */
    [data-testid="stSidebar"] {display: none;} /* Ẩn sidebar */
    
    /* 2. CĂN CHỈNH LẠI BỐ CỤC CHO ĐẸP */
    /* Đẩy nội dung lên trên cùng vì đã ẩn header */
    .block-container {
        padding-top: 2rem !important; 
        padding-bottom: 60px !important;
    }

    /* 3. GIAO DIỆN CHUNG */
    .stApp {
        background-color: #ffffff;
        font-family: 'Segoe UI', Helvetica, Arial, sans-serif;
    }

    /* Header Banner Chuyên nghiệp */
    .main-header {
        background: linear-gradient(90deg, #ce181e 0%, #003b8e 100%); /* Đỏ PCCC sang Xanh CA */
        padding: 2rem 1rem;
        border-radius: 15px;
        box-shadow: 0 4px 15px rgba(0,0,0,0.15);
        margin-bottom: 2rem;
        text-align: center;
        color: white;
    }
    
    /* Footer Bản Quyền Riêng */
    .custom-footer {
        position: fixed;
        left: 0;
        bottom: 0;
        width: 100%;
        background-color: #f8f9fa;
        color: #555;
        text-align: center;
        padding: 10px;
        font-size: 0.8rem;
        border-top: 1px solid #e0e0e0;
        z-index: 9999;
    }

    /* Tinh chỉnh Chat Input */
    .stChatInput textarea {
        border-radius: 30px !important;
        padding: 12px 20px !important;
        border: 1px solid #dfe1e5 !important;
        box-shadow: 0 2px 5px rgba(0,0,0,0.05);
    }
    .stChatInput textarea:focus {
        border-color: #ce181e !important;
        box-shadow: 0 2px 8px rgba(206, 24, 30, 0.2);
    }

    /* Chat Message Styling */
    [data-testid="stChatMessage"] {
        padding: 1rem;
        border-radius: 10px;
        margin-bottom: 10px;
        background-color: transparent;
    }
    
    /* User Message Background */
    [data-testid="stChatMessage"][data-testid="user"] {
        background-color: #f1f3f4;
    }

    /* Content Typography */
    .response-content {
        line-height: 1.6;
        color: #1a1a1a;
        font-size: 1rem;
    }
    .response-content strong {
        color: #ce181e; 
    }
</style>
""", unsafe_allow_html=True)

# --- 3. HEADER GIAO DIỆN ---
st.markdown("""
<div class="main-header">
    <div style="font-size: 1.1rem; font-weight: 600; margin-bottom: 5px; opacity: 0.95; letter-spacing: 0.5px;">CÔNG AN TỈNH PHÚ THỌ</div>
    <div style="font-size: 1.6rem; font-weight: 800; margin-bottom: 15px; line-height: 1.2; text-transform: uppercase;">PHÒNG CẢNH SÁT PCCC VÀ CNCH</div>
    <div style="font-size: 1.4rem; font-weight: 700; margin-bottom: 10px; color: #ffe082; text-transform: uppercase; text-shadow: 1px 1px 2px rgba(0,0,0,0.2);">TRỢ LÝ AI VỀ PCCC VÀ CNCH</div>
    <div style="font-size: 0.95rem; font-weight: 300; font-style: italic; opacity: 0.9;">(Được phát triển bởi Đại úy Phạm Tùng Linh)</div>
</div>
""", unsafe_allow_html=True)

# --- 4. FOOTER BẢN QUYỀN ---
st.markdown("""
<div class="custom-footer">
    © 2025 Bản quyền thuộc về Đại úy Phạm Tùng Linh - PC07 Công an tỉnh Phú Thọ
</div>
""", unsafe_allow_html=True)

# --- 5. KẾT NỐI API (FAILOVER) ---
API_KEYS_LIST = []
try:
    if "GEMINI_API_KEYS" in st.secrets: 
        keys_string = st.secrets["GEMINI_API_KEYS"]
        API_KEYS_LIST = [k.strip() for k in keys_string.split(",") if k.strip()]
    elif "GEMINI_API_KEY" in st.secrets:
        API_KEYS_LIST = [st.secrets["GEMINI_API_KEY"]]
    
    if not API_KEYS_LIST: st.error("❌ Lỗi: Thiếu API Key hệ thống."); st.stop()
    DRIVE_FOLDER_ID = st.secrets["DRIVE_FOLDER_ID"]
    GCP_JSON = json.loads(st.secrets["GCP_JSON"])
except Exception as e: st.error(f"⚠️ Lỗi cấu hình: {str(e)}"); st.stop()

# --- 6. BỘ NÃO THAM MƯU (ROUTER) ---
ROUTER_INSTRUCTION = """
Bạn là Tham mưu trưởng PCCC xuất sắc. NHIỆM VỤ TỐI THƯỢNG: Phân tích sâu ngữ nghĩa, phán đoán chính xác "Ý định thực sự" (Intent) của người dùng thông qua câu hỏi đời thường, từ đó chọn ĐÚNG và ĐỦ các tài liệu pháp lý tương ứng. 

QUY TRÌNH TƯ DUY BẮT BUỘC:
- Bước 1: Dịch ngôn ngữ đời thường sang thuật ngữ pháp lý. (Ví dụ: "xin giấy cháy nổ" = Thẩm duyệt/Nghiệm thu; "đền bù cháy", "mua bảo hiểm" = Bảo hiểm cháy nổ bắt buộc; "mấy cửa ra" = Lối thoát nạn; "công an phạt hay ủy ban phạt" = Thẩm quyền xử lý vi phạm).
- Bước 2: Tự hỏi "Bản chất cốt lõi của câu hỏi này thuộc lĩnh vực quản lý nhà nước nào?".
- Bước 3: Áp chiếu vào các Giỏ tài liệu dưới đây để bốc đúng file.

DANH SÁCH CÁC GIỎ TÀI LIỆU VÀ BẢN CHẤT CỦA CHÚNG:
1. GIỎ PHÂN CẤP QUẢN LÝ (THẨM QUYỀN VÀ DANH MỤC):
   - Bản chất: Xác định cơ sở này thuộc diện nào, do cấp nào quản lý (Công an PC07, Công an huyện, hay UBND cấp xã), tra cứu các Phụ lục phân loại.
   - Hành động: BẮT BUỘC CHỌN [Nghị định 105].

2. GIỎ THỦ TỤC HÀNH CHÍNH & PHÁP LÝ CHUNG (HỒ SƠ, BÁO CÁO, BẢO HIỂM):
   - Bản chất: Các vấn đề trên giấy tờ, quy trình làm việc với cơ quan nhà nước và TỔ CHỨC LỰC LƯỢNG. Bao gồm: Điều kiện an toàn, hồ sơ thiết kế, nghiệm thu, kiểm tra định kỳ, trách nhiệm chủ cơ sở, trách nhiệm chủ đầu tư, trách nhiệm chủ phương tiện, phương án chữa cháy, phương án cứu nạn cứu hộ, huấn luyện nghiệp vụ, thành lập ĐỘI PCCC CƠ SỞ, lực lượng dân phòng, chuyên ngành, người được phân công nhiệm vụ PCCC, BẢO HIỂM CHÁY NỔ BẮT BUỘC, BÁO CÁO định kỳ và các loại biểu mẫu.
   - Hành động: BẮT BUỘC CHỌN [Luật PCCC và CNCH], [Nghị định 105], [Thông tư 36].
   - CẤM: Hỏi về "phương án chữa cháy" của CƠ SỞ thì TUYỆT ĐỐI KHÔNG chọn Thông tư 37.

3. GIỎ XỬ PHẠT (CHẾ TÀI VI PHẠM):
   - Bản chất: Người dùng hỏi về hành vi sai phạm, bị phạt bao nhiêu tiền, chức danh nào có quyền ký quyết định phạt, tước giấy phép.
   - Hành động: BẮT BUỘC CHỌN [Nghị định 106], [Nghị định 189].

4. GIỎ CƯỠNG CHẾ (KHÔNG NỘP PHẠT):
   - Bản chất: Áp dụng khi đã có quyết định xử phạt nhưng người vi phạm chây ỳ, nộp muộn, không nộp phạt. Cần các biện pháp cưỡng chế thu tiền, kê biên tài sản, khấu trừ lương.
   - Hành động: BẮT BUỘC CHỌN [Nghị định 296].

5. GIỎ KỸ THUẬT - KIẾN TRÚC & XÂY DỰNG (QCVN 06):
   - Bản chất: Các yếu tố "cứng" gắn liền với vỏ/khung công trình: Đường giao thông cho xe cứu hỏa, khoảng cách an toàn, bậc chịu lửa, lối thoát nạn (cửa, cầu thang, hành lang), ngăn cháy lan, thông gió, hút khói.
   - Hành động: BẮT BUỘC CHỌN [QCVN 06].

6. GIỎ KỸ THUẬT - LẮP ĐẶT THIẾT BỊ PCCC (QCVN 10):
   - Bản chất: Các yếu tố "mềm" lắp thêm vào công trình: Cảm biến báo cháy, bình chữa cháy, đầu phun Sprinkler, máy bơm, bể nước, họng nước vách tường, trụ cấp nước.
   - Hành động: BẮT BUỘC CHỌN [QCVN 10].

7. GIỎ CHIẾN THUẬT & QUÂN ĐỘI:
   - Bản chất: Nghiệp vụ thực chiến của Cảnh sát PCCC khi ra trận: Chỉ huy, chiến thuật dập lửa, phối hợp quân đội, dân quân.
   - Hành động: CHỌN [Thông tư 37], [Luật PCCC], các file chứa từ [QUÂN ĐỘI], [CV HD].

OUTPUT: CHỈ trả về danh sách tên file chính xác có trong kho. Ngăn cách bằng dấu phẩy. TUYỆT ĐỐI KHÔNG in ra quá trình suy luận để hệ thống tải file không bị lỗi.
"""

def smart_router(user_query, available_files):
    file_list_str = ", ".join(available_files)
    prompt = f"""{ROUTER_INSTRUCTION}\n\nDANH SÁCH FILE: {file_list_str}\n\nCÂU HỎI: "{user_query}"\n\nCHỌN TÀI LIỆU:"""
    for key in API_KEYS_LIST:
        try:
            genai.configure(api_key=key)
            model = genai.GenerativeModel('gemini-1.5-flash')
            response = model.generate_content(prompt)
            return response.text.strip()
        except: continue
    return ""

# --- 7. BỘ NÃO CHUYÊN GIA (EXPERT) ---
SYSTEM_PROMPT_EXPERT = """
VAI TRÒ: Trợ lý AI về PCCC và CNCH - Phòng PC07 Phú Thọ.

🛑 NGUYÊN TẮC CỐT TỬ:
1. Trả lời ngắn gọn, đúng trọng tâm, văn phong hành chính chuyên nghiệp.
2. Tuyệt đối không sáng tạo ngoài văn bản.
3. TUYỆT ĐỐI KHÔNG sử dụng kiến thức có sẵn trên mạng (như NĐ 136 cũ hay Luật cũ). CHỈ ĐƯỢC PHÉP lấy thông tin và căn cứ từ văn bản được cung cấp.
4. TUYỆT ĐỐI KHÔNG để lộ các từ khóa quy trình như "RULE 1", "RULE 2", "BƯỚC 1", "GIỎ"... vào trong câu trả lời. Hệ thống phải suy luận ngầm và chỉ xuất ra kết quả cuối cùng tự nhiên nhất.

🔴 RULE 1: XÁC ĐỊNH THẨM QUYỀN QUẢN LÝ (QUAN TRỌNG - THEO NĐ 105/2025):
   BẮT BUỘC thực hiện đúng 2 BƯỚC sau:
   - BƯỚC 1: ĐỐI CHIẾU PHỤ LỤC I và PHỤ LỤC II (Nghị định 105/2025/NĐ-CP).
     + So sánh các chỉ số: Số tầng, Khối tích, Diện tích với Phụ lục I và Phụ lục II.
   - BƯỚC 2: KẾT LUẬN (QUY TẮC ƯU TIÊN TUYỆT ĐỐI):
     + Nếu cơ sở đạt tiêu chí Phụ lục II -> PHÒNG CẢNH SÁT PCCC & CNCH (PC07) quản lý.
     + Lưu ý đặc biệt: Dù diện tích nhỏ (thuộc Phụ lục I) nhưng Số tầng cao (thuộc Phụ lục II) -> Vẫn là PC07 quản lý.
     + Chỉ khi nào KHÔNG đạt Phụ lục II mà CHỈ đạt Phụ lục I -> Mới do UBND CẤP XÃ quản lý.

🔴 RULE 2: XỬ LÝ / XỬ PHẠT VI PHẠM (NĐ 106 + 189):
   - KHI NGƯỜI DÙNG HỎI: "Xử lý như nào", "Bị sao", "Phạt bao nhiêu", "Lỗi này thế nào"... -> HIỂU NGAY LÀ HỎI VỀ XỬ PHẠT HÀNH CHÍNH.
   - ⚠️ ĐỒNG NHẤT NGÔN NGỮ: "chưa" = "không" (VD: "chưa huấn luyện" = "không huấn luyện", "chưa thẩm duyệt" = "không thẩm duyệt"). Trợ lý BẮT BUỘC hiểu đồng nhất để quét trúng hành vi.
   - ⚠️ ĐỊNH DẠNG VÀ TƯ DUY BẮT BUỘC (Trình bày chính xác theo template, in đậm tiêu đề, xuống dòng rõ ràng):

     **1. HÀNH VI:** [Tên hành vi chính xác trong NĐ 106]

     **2. MỨC PHẠT TIỀN:**
     - Cá nhân: ... (Căn cứ: Điểm... Khoản... Điều... NĐ 106).
     - Tổ chức: ... (Gấp 2 lần mức cá nhân).

     **3. HÌNH THỨC PHẠT BỔ SUNG & KHẮC PHỤC HẬU QUẢ (KPHQ):**
     - Phạt bổ sung: [Có/Không] -> Nêu rõ TÊN biện pháp (Căn cứ NĐ 106).
     - Biện pháp KPHQ: [Có/Không] -> Nêu rõ TÊN biện pháp (VD: Buộc tổ chức huấn luyện, Buộc tháo dỡ...) (Căn cứ NĐ 106).

     **4. THẨM QUYỀN XỬ PHẠT (ĐỐI CHIẾU KÉP CHUẨN XÁC THEO NĐ 189):**
     * CHỈ XÉT 6 chức danh: Chiến sĩ CA, Đội trưởng, Trưởng CA cấp xã, Trưởng Phòng PC07, Giám đốc CA cấp tỉnh, Chủ tịch UBND cấp tỉnh. (TUYỆT ĐỐI KHÔNG CÓ Đội trưởng cấp huyện).
     * BẮT BUỘC THỰC HIỆN BƯỚC LỌC KÉP SAU VỚI TỪNG CHỨC DANH (Dựa trên NĐ 189/2025/NĐ-CP):
       - ĐIỀU KIỆN 1 (TIỀN): Thẩm quyền phạt tiền tối đa của chức danh phải >= Mức phạt tiền của hành vi (Lưu ý phân biệt mức cá nhân/tổ chức).
       - ĐIỀU KIỆN 2 (PHẠT BỔ SUNG & KPHQ): ĐỌC KỸ quy định thẩm quyền của chức danh đó trong NĐ 189. Nếu hành vi ở Mục 3 có Phạt bổ sung hoặc KPHQ, BẮT BUỘC chức danh đó phải CÓ QUYỀN áp dụng ĐÚNG LOẠI Phạt bổ sung/KPHQ đó. (Ví dụ: Nếu Mục 3 yêu cầu "Buộc tổ chức huấn luyện", AI phải kiểm tra xem Đội trưởng, Trưởng CA xã... có được giao quyền áp dụng biện pháp "Buộc tổ chức huấn luyện" theo NĐ 189 không. Nếu KHÔNG -> LOẠI NGAY LẬP TỨC chức danh đó, bất kể mức tiền thỏa mãn).
     [CHỈ liệt kê bằng gạch đầu dòng những người VƯỢT QUA CẢ 2 ĐIỀU KIỆN trên]:
     - [Tên chức danh 1]
     - [Tên chức danh 2]

     **5. KIẾN NGHỊ:**
     Trình [Tên chức danh cấp xã thấp nhất CÒN LẠI TRONG DANH SÁCH MỤC 4] và [Tên chức danh cấp tỉnh thấp nhất CÒN LẠI TRONG DANH SÁCH MỤC 4: Đội trưởng hoặc Trưởng Phòng PC07 hoặc Giám đốc Công an tỉnh hoặc Chủ tịch UBND tỉnh] ký quyết định. (TUYỆT ĐỐI KHÔNG kiến nghị chức danh đã bị loại ở Mục 4).
  
🔴 RULE 3: CƯỠNG CHẾ / KHÔNG NỘP PHẠT (NĐ 296/2025):
   - Khi hỏi về việc không nộp tiền, nộp chậm, chây ỳ -> Dùng NĐ 296/2025/NĐ-CP.
   - Trả lời các biện pháp: Khấu trừ lương/thu nhập, Khấu trừ tiền từ tài khoản, Kê biên tài sản...

🔴 RULE 4: TRÁCH NHIỆM / ĐIỀU KIỆN / HỒ SƠ / KIỂM TRA / NGHIỆM THU / THẨM ĐỊNH / PHÒNG CHÁY / BẢO VỆ HIỆN TRƯỜNG/ PHƯƠNG ÁN CHỮA CHÁY:
   # NGUYÊN TẮC TRA CỨU THEO THỨ BẬC PHÁP LÝ (HIERARCHICAL CASCADING)
   Khi nhận được bất kỳ câu hỏi nào liên quan đến các chủ đề trên, bạn BẮT BUỘC phải thực hiện luồng tra cứu tuần tự sau đây. Tuyệt đối KHÔNG được dừng lại hoặc từ chối giữa chừng nếu chưa quét hết 3 cấp độ:
   - BƯỚC 1 (QUÉT LUẬT): Ưu tiên tìm kiếm trong "Luật PCCC và CNCH". Nếu Luật có quy định -> Trích dẫn ngay. 
   - BƯỚC 2 (CHUYỂN TIẾP XUỐNG NGHỊ ĐỊNH): Nếu Luật không quy định chi tiết (đặc biệt là các câu hỏi về Biểu mẫu, Hồ sơ, Thẩm quyền phê duyệt cụ thể) -> TỰ ĐỘNG bỏ qua Luật và quét toàn diện vào Nghị định (VD: Nghị định 105), bao gồm cả phần Phụ lục. Nếu có -> Trích dẫn nguyên văn.
   - BƯỚC 3 (CHUYỂN TIẾP XUỐNG THÔNG TƯ): Nếu Nghị định tiếp tục không có, hoặc có điều khoản ghi "thực hiện theo hướng dẫn của Bộ Công an" -> TỰ ĐỘNG quét tiếp xuống các Thông tư (VD: Thông tư 36, Thông tư 37), bao gồm cả Phụ lục. Nếu có -> Trích dẫn.
   - BƯỚC 4 (CHỐT CHẶN CUỐI CÙNG): Bạn CHỈ ĐƯỢC PHÉP trả lời từ chối (theo nguyên tắc số 7) SAU KHI đã quét cạn kiệt cả 3 cấp độ (Luật -> Nghị định -> Thông tư) từ các Điều khoản đầu tiên cho đến Phụ lục biểu mẫu cuối cùng mà vẫn không có kết quả.
   
    
🟢 RULE 5: CÁC LĨNH VỰC KHÁC VÀ TRÌNH BÀY QCVN 06, QCVN 10:
   - Kỹ thuật: BẮT BUỘC tra cứu số liệu cụ thể từ QCVN 06:2022/BXD (hoặc sửa đổi) và QCVN 10:2025/BCA.

   - ⚠️ YÊU CẦU TRÌNH BÀY ĐỐI VỚI QCVN 06:2022/BXD:
     Khi trả lời QCVN 06, BẮT BUỘC: 1. Trích dẫn ĐẦY ĐỦ nguyên văn nội dung. 2. Ghi CHÍNH XÁC Mục/Điều/Bảng. Không được tóm tắt.

   - ⚠️ ĐỊNH DẠNG BẮT BUỘC ĐỐI VỚI QCVN 10:2025/BCA (CẤM VIẾT THÀNH ĐOẠN VĂN):
     Mọi hệ thống/phương tiện BẮT BUỘC phải trình bày theo đúng 3 dòng sau, không được sai lệch:
     [Tên hệ thống/phương tiện]:
     - Yêu cầu: [Chỉ ghi "Phải trang bị" HOẶC "Không thuộc diện phải trang bị"]
     - Căn cứ: [Trích dẫn rõ ràng Bảng, Mục tương ứng. Ghi rõ số liệu điều kiện nếu có]

   - ⚠️ LỆNH CHỐNG ẢO GIÁC ĐỐI VỚI "HỆ THỐNG CẤP NƯỚC CHỮA CHÁY NGOÀI NHÀ" (BẢNG C.1):
     + LƯU Ý TỐI QUAN TRỌNG: Bảng C.1 TUYỆT ĐỐI KHÔNG CÓ các loại hình như "Nhà nghỉ", "Khách sạn", "Karaoke", "Nhà ở riêng lẻ", "Cơ sở lưu trú". AI cấm được nhầm lẫn Bảng C.1 với Bảng A.1 và B.1.
     + TRƯỜNG HỢP 1 (CƠ SỞ KHÔNG CÓ TRONG BẢNG C.1 - Ví dụ: Nhà nghỉ, Khách sạn...):
       Hệ thống cấp nước chữa cháy ngoài nhà:
       - Yêu cầu: Không thuộc diện phải trang bị.
       - Căn cứ: Loại hình cơ sở này không nằm trong 10 mục yêu cầu phải trang bị tại Bảng C.1 Phụ lục C QCVN 10:2025/BCA. (CẤM NHẮC ĐẾN MỤC 2.3.2 Ở TRƯỜNG HỢP NÀY).
     + TRƯỜNG HỢP 2 (CÓ TÊN ĐÚNG TRONG BẢNG C.1 VÀ ĐẠT QUY MÔ):
       Hệ thống cấp nước chữa cháy ngoài nhà:
       - Yêu cầu: Phải trang bị.
       - Căn cứ: [Trích đúng số thứ tự Mục trong Bảng C.1]. Lưu ý: Theo Mục 2.3.2 QCVN 10:2025/BCA, cho phép không trang bị khi nhà cách trụ/bến lấy nước chữa cháy dưới 400m...
       
   - Chữa cháy, chỉ huy chữa cháy: Căn cứ Luật PCCC, Nghị định 105 và Thông tư 37.
   - Quân đội: Căn cứ CV Hướng dẫn phối hợp.
"""

# Đã bổ sung biến timeout_val vào hàm để nhận lệnh thời gian từ bên dưới
def call_gemini_expert_exhaustive(prompt, context, timeout_val=60, image_obj=None):
    # Đã trả lại đầy đủ đội hình Model siêu mạnh theo đúng nguyên bản của Đại úy
    TARGET_MODELS = ["gemini-2.5-flash", "gemini-2.0-flash", "gemini-2.0-flash-exp", "gemini-1.5-pro", "gemini-1.5-flash"]
    
    # LUÔN ÉP KHUNG SYSTEM_PROMPT KỂ CẢ KHI TÀI LIỆU RỖNG (CHỐNG ẢO GIÁC LUẬT CŨ)
    if not context: 
        full_prompt = f"{SYSTEM_PROMPT_EXPERT}\n\n=== TÀI LIỆU HỖ TRỢ ===\n(Hiện tại không tìm thấy tài liệu nào khớp với câu hỏi)\n\n=== CÂU HỎI ===\n{prompt}\n\nLƯU Ý CỐT TỬ: Nếu đây là câu chào hỏi, hãy chào lại. Nếu là câu hỏi nghiệp vụ PCCC, BẮT BUỘC trả lời đúng nguyên văn: 'Xin lỗi đồng chí, tôi chưa tìm thấy quy định cụ thể về vấn đề này trong hệ thống tài liệu hiện tại của PC07 Phú Thọ.' TUYỆT ĐỐI KHÔNG sử dụng kiến thức trên mạng để trả lời."
    else: 
        full_prompt = f"{SYSTEM_PROMPT_EXPERT}\n\n=== TÀI LIỆU HỖ TRỢ ===\n{context}\n\n=== LƯU Ý TỐI QUAN TRỌNG KHI CÓ ẢNH (BỘ LỌC MẮT THẦN) ===\nBƯỚC 1 - NHẬN DIỆN ẢNH: Bắt buộc kiểm tra xem bức ảnh có chứa bản vẽ kỹ thuật, mặt bằng kiến trúc, công trình, lối thoát nạn, thiết bị PCCC, hoặc hiện trường sự cố hay không.\nBƯỚC 2 - TỪ CHỐI (NẾU ẢNH KHÔNG LIÊN QUAN): Nếu người dùng tải lên ảnh chó mèo, đồ ăn, phong cảnh, ảnh selfie người, hoặc các hình ảnh vớ vẩn không thuộc lĩnh vực xây dựng/PCCC... AI BẮT BUỘC chỉ được trả lời: 'Bức ảnh này dường như không liên quan đến nghiệp vụ PCCC & CNCH. Đồng chí vui lòng tải lên bản vẽ thiết kế hoặc ảnh chụp thực tế hiện trường để tôi hỗ trợ phân tích.' -> VÀ DỪNG LẠI NGAY LẬP TỨC, không phân tích gì thêm.\nBƯỚC 3 - PHÂN TÍCH CHUYÊN SÂU (NẾU ẢNH HỢP LỆ): Đóng vai trò Cán bộ Thẩm duyệt/Kiểm tra PC07. Hãy quét bằng mắt để phân tích cụ thể: 1. Chiều mở cửa thoát nạn. 2. Số lượng và vị trí thang bộ. 3. Các hành lang cụt. 4. Bố trí phương tiện (bình chữa cháy, báo cháy). Chỉ ra rõ các điểm vi phạm QCVN 06:2022/BXD và QCVN 10:2025/BCA dựa trên hình học không gian và logic thực tế.\n\n=== CÂU HỎI ===\n{prompt}"
    
    payload = [full_prompt]
    if image_obj:
        # Ép chuẩn màu RGB để API đọc ảnh mượt mà, không báo lỗi
        payload.append(image_obj.convert('RGB'))
    
    last_error = ""
    for model_name in TARGET_MODELS:
        for index, key in enumerate(API_KEYS_LIST):
            try:
                genai.configure(api_key=key)
                model = genai.GenerativeModel(model_name)
                response = model.generate_content(payload, request_options={'timeout': timeout_val})
                return response.text
            except Exception as e:
                last_error = str(e)
                continue 
    return f"⚠️ Hệ thống đang bận. Vui lòng thử lại. (Error: {last_error})"

# --- 8. NẠP DỮ LIỆU ---
@st.cache_data(ttl=7200, show_spinner=False)
def load_database_final():
    if not GCP_JSON or not DRIVE_FOLDER_ID: return {}, []
    try:
        creds = service_account.Credentials.from_service_account_info(GCP_JSON)
        service = build('drive', 'v3', credentials=creds)
        db = {} 
        logs = []
        processed = set()
        
        # Keywords bao gồm Nghị định 105, 106, 296...
        keywords = ["189", "106", "105", "296", "36", "37", "48", "luat", "huy dong", "quan doi", "du thao", "phoi hop", "cv hd", "doi 3", "qcvn", "10:2025", "06", "10"]
        files = []
        for k in keywords:
            try: files.extend(service.files().list(q=f"'{DRIVE_FOLDER_ID}' in parents and trashed=false and name contains '{k}'", fields="files(id, name)").execute().get('files', []))
            except: pass
        
        # Fallback lấy thêm file nếu ít
        if len(files) < 5:
             try: files.extend(service.files().list(q=f"'{DRIVE_FOLDER_ID}' in parents and trashed=false", pageSize=50, fields="files(id, name)").execute().get('files', []))
             except: pass

        for f in files:
            if f['id'] in processed: continue
            processed.add(f['id'])
            if "144" in f['name'] and "106" not in f['name']: continue

            try:
                fh = io.BytesIO()
                downloader = MediaIoBaseDownload(fh, service.files().get_media(fileId=f['id']))
                done = False
                while not done: _, done = downloader.next_chunk()
                fh.seek(0)
                
                text = ""
                if f['name'].endswith(".docx"):
                    doc = Document(fh)
                    text = "\n".join([p.text for p in doc.paragraphs])
                    for t in doc.tables:
                        for r in t.rows: text += " | ".join([c.text.strip() for c in r.cells]) + "\n"
                elif f['name'].endswith(".pdf"):
                    reader = PdfReader(fh)
                    text = "\n".join([p.extract_text() for p in reader.pages if p.extract_text()])
                
                if text:
                    db[f['name']] = text
                    logs.append(f"✅ {f['name']}")
            except: continue
        return db, logs
    except Exception as e: return None, []

# --- 9. KHỞI ĐỘNG ---
with st.spinner('🔄 Đang khởi tạo hệ thống nghiệp vụ...'):
    database, logs = load_database_final()

if not database: st.error("❌ Không thể kết nối Kho dữ liệu. Vui lòng kiểm tra lại cấu hình."); st.stop()
# --- THỦ THUẬT TỰ ĐỘNG CUỘN TRANG (AUTO-SCROLL) ---
def scroll_to_bottom():
    st.components.v1.html(
        """
        <script>
            var body = window.parent.document.querySelector('.main');
            if (body) {
                body.scrollTop = body.scrollHeight;
            }
        </script>
        """,
        height=0, width=0
    )

# --- 10. CHAT LOGIC ---
if "messages" not in st.session_state: 
    # KHỞI TẠO LỜI CHÀO BAN ĐẦU
    st.session_state.messages = [{
        "role": "assistant", 
        "content": "Xin chào! Tôi là trợ lý AI về PCCC và CNCH do Đại úy Phạm Tùng Linh - Phòng PC07 phát triển. Hãy đặt câu hỏi để tôi trả lời."
    }]

for m in st.session_state.messages:
    with st.chat_message(m["role"], avatar="👮‍♂️" if m["role"] == "user" else "🔥"): 
        st.markdown(f'<div class="response-content">{m["content"]}</div>', unsafe_allow_html=True)

# --- TẠO CÁC NÚT CÂU HỎI MẪU ---
if "suggested_prompt" not in st.session_state:
    st.session_state.suggested_prompt = None

st.markdown("<p style='font-size: 0.9rem; color: #555;'>💡 <b>Gợi ý tra cứu nhanh:</b></p>", unsafe_allow_html=True)

# Khai báo ĐÚNG 4 CỘT
col1, col2, col3, col4 = st.columns(4)

with col1:
    if st.button("🏢 Phương tiện PCCC QC10"): 
        st.session_state.suggested_prompt = "Nhà nghỉ 5 tầng cần trang bị phương tiện PCCC gì?"
with col2:
    if st.button("🚪 Lối thoát nạn QC06"): 
        st.session_state.suggested_prompt = "Quy định về lối thoát nạn như thế nào?"
with col3:
    if st.button("💰 Xử phạt vi phạm hành chính"): 
        st.session_state.suggested_prompt = "Cơ sở sử dụng người chưa được huấn luyện PCCC bị xử lý như nào?"
with col4:
    if st.button("📋 Quy định pháp luật PCCC"): 
        st.session_state.suggested_prompt = "Cơ sở karaoke cao 5 tầng do ai quản lý?"

# --- NHẬN LỆNH TỪ CHAT HOẶC TỪ NÚT BẤM ---
# LƯU Ý: ĐÂY LÀ CHỖ DUY NHẤT CÓ st.chat_input ĐỂ TRÁNH LỖI TRÙNG LẶP
prompt = st.chat_input("Nhập nội dung cần tra cứu...") or st.session_state.suggested_prompt

if prompt:
    st.session_state.suggested_prompt = None # Reset lại nút bấm sau khi gửi
    
    prompt_lower = prompt.lower()
    
    st.session_state.messages.append({"role": "user", "content": prompt})
    st.chat_message("user", avatar="👮‍♂️").write(prompt)
    
    # ======== GẮN LẠI LẦN 1: CUỘN XUỐNG KHI USER VỪA GỬI CÂU HỎI ========
    scroll_to_bottom()
    
    with st.chat_message("assistant", avatar="🔥"):
        with st.spinner("🧠 Đang suy nghĩ, bạn chờ chút, mình trả lời ngay đây..."):
            
            # Router
            all_files = list(database.keys())
            selected_files_str = smart_router(prompt, all_files)
            
            # Retrieve
            relevant_context = ""
            if selected_files_str:
                for fname in all_files:
                    if fname in selected_files_str: relevant_context += f"--- VĂN BẢN: {fname} ---\n{database[fname]}\n"
            
            # Backup Retrieve (BỔ SUNG LOGIC: Cưỡng chế & Xử lý & Phương án)
            if not relevant_context:
                for fname, content in database.items():
                    is_enforcement = ("cưỡng chế" in prompt_lower or "không nộp" in prompt_lower or "chậm nộp" in prompt_lower or "chây ỳ" in prompt_lower)
                    is_penalty = ("phạt" in prompt_lower or "lỗi" in prompt_lower or "xử lý" in prompt_lower)
                    is_military = ("quân đội" in prompt_lower or "chi viện" in prompt_lower)
                    
                    is_tech_06 = any(kw in prompt_lower for kw in ["khoảng cách", "ngăn cháy", "thông gió", "hút khói", "chống cháy lan", "đường lối", "bãi đỗ", "vật liệu", "kích thước", "thoát nạn", "lối vào", "06", "qc06"]) and not is_penalty
                    is_tech_10 = any(kw in prompt_lower for kw in ["trang bị", "lắp đặt", "hệ thống", "10", "qc10"]) and not is_penalty
                    
                    is_manage = ("trách nhiệm" in prompt_lower or "hồ sơ" in prompt_lower or "quản lý" in prompt_lower or "điều kiện" in prompt_lower or "kiểm tra" in prompt_lower or "phương án" in prompt_lower or "mẫu" in prompt_lower)
                    
                    # Nếu hỏi phương án/mẫu thì tuyệt đối KHÔNG phải là TT37
                    is_force = ("lực lượng" in prompt_lower or "chữa cháy" in prompt_lower) and not ("phương án" in prompt_lower or "mẫu" in prompt_lower)

                    if is_enforcement and "296" in fname: 
                        relevant_context += "--- " + fname + " ---\n" + content + "\n"
                    elif is_military and any(x in fname.lower() for x in ["quan doi", "du thao", "phoi hop", "cv hd", "doi 3"]):
                        relevant_context += "--- " + fname + " ---\n" + content + "\n"
                    elif is_penalty and any(x in fname for x in ["106", "189"]):
                        relevant_context += "--- " + fname + " ---\n" + content + "\n"
                    elif is_tech_06 and "06" in fname: 
                        relevant_context += "--- " + fname + " ---\n" + content + "\n"
                    elif is_tech_10 and "10" in fname: 
                        relevant_context += "--- " + fname + " ---\n" + content + "\n"
                    elif is_manage and any(x in fname for x in ["luat", "105", "36", "136", "50"]):
                        relevant_context += "--- " + fname + " ---\n" + content + "\n"
                    elif is_force and "37" in fname:
                        relevant_context += "--- " + fname + " ---\n" + content + "\n"

            # TÍNH TOÁN THỜI GIAN CHỜ THEO LOẠI TÀI LIỆU
            current_timeout = 60
            if "06" in prompt_lower or "qcvn" in prompt_lower or "kỹ thuật" in prompt_lower or "lối thoát" in prompt_lower or "thoát nạn" in prompt_lower or "ngăn cháy" in prompt_lower:
                current_timeout = 180
            elif "06" in relevant_context.lower() or "qcvn" in relevant_context.lower():
                current_timeout = 180

            # Generate
            response_text = call_gemini_expert_exhaustive(prompt, relevant_context, current_timeout)
        
        st.markdown(f'<div class="response-content">{response_text}</div>', unsafe_allow_html=True)
        st.session_state.messages.append({"role": "assistant", "content": response_text})
        
    # ======== GẮN LẠI LẦN 2: CUỘN XUỐNG KHI AI TRẢ LỜI XONG ========
    scroll_to_bottom()
